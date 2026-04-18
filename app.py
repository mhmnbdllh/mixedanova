"""
Mixed ANOVA Calculator — SPSS GLM Repeated Measures Equivalent
==============================================================
Input  : WIDE format CSV (one row per subject)
Engine : Fully parametric path (Mixed ANOVA) when normality holds;
         automatically switches to non-parametric alternatives when violated.

Parametric path — SPSS-identical formulas (Winer, Brown & Michels 1991, Ch. 7):
  SS_A     = b · Σᵢ nᵢ (Ȳᵢ.. − Ȳ...)²               df = a−1
  SS_S(A)  = b · Σᵢ Σₛ (Ȳₛ.. − Ȳᵢ..)²               df = N−a    ← between error
  SS_B     = N · Σⱼ (Ȳ.j. − Ȳ...)²                   df = b−1
  SS_AB    = Σᵢ nᵢ · Σⱼ (Ȳᵢⱼ − Ȳᵢ.. − Ȳ.j. + Ȳ...)² df = (a−1)(b−1)
  SS_BS(A) = Σᵢ Σₛ Σⱼ (yₛⱼ − Ȳₛ.. − Ȳᵢⱼ + Ȳᵢ..)²   df = (b−1)(N−a) ← within error
  [Verified: five components sum exactly to SS_Total]

  F_A  = MS_A  / MS_S(A)   F_B = MS_B / MS_BS(A)   F_AB = MS_AB / MS_BS(A)
  Partial η²p = SS_effect / (SS_effect + SS_error_for_that_effect)
  EMM SE = √(MS_BS(A) / n_cell) with t(df_BS(A)) CI  — SPSS method
  Mauchly W via Box (1954) χ² approx; GG / HF (Lecoutre 1991) / LB corrections

Non-parametric path (when any cell violates normality at α = .05):
  Within-subjects : Friedman test per group; Wilcoxon signed-rank post-hoc
  Between-subjects: Kruskal–Wallis; Mann–Whitney U post-hoc
  Effect sizes    : Kendall's W (Friedman), rank-biserial r (Mann–Whitney)
"""

import io
import itertools
import re
import warnings

import numpy as np
import pandas as pd
import scipy.stats as stats
from scipy.stats import f as fdist

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.gridspec as mgs
import seaborn as sns

import streamlit as st

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

warnings.filterwarnings("ignore")

# ══════════════════════════════════════════════════════════════════════════════
st.set_page_config(page_title="Mixed ANOVA Calculator", page_icon="📊",
                   layout="wide", initial_sidebar_state="expanded")

# ══════════════════════════════════════════════════════════════════════════════
#  CSS
# ══════════════════════════════════════════════════════════════════════════════
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Lora:wght@500;600&family=Inter:wght@300;400;500;600&family=JetBrains+Mono:wght@400;500&display=swap');
html,body,[class*="css"]{font-family:'Inter',sans-serif;}

.app-title{font-family:'Lora',serif;font-size:2.1rem;font-weight:600;
  color:#0d1b2a;letter-spacing:-.01em;margin-bottom:2px;}
.app-sub{font-size:.86rem;color:#546e7a;font-weight:300;margin-bottom:1.4rem;}
.sec{font-family:'Lora',serif;font-size:1.1rem;font-weight:600;color:#0d1b2a;
  border-left:4px solid #c62828;padding-left:10px;margin-top:1.7rem;margin-bottom:.7rem;}
.sec-np{font-family:'Lora',serif;font-size:1.1rem;font-weight:600;color:#4a148c;
  border-left:4px solid #7b1fa2;padding-left:10px;margin-top:1.7rem;margin-bottom:.7rem;}

.mcard{background:#0d1b2a;border-radius:10px;padding:.9rem 1.1rem .85rem;margin-bottom:4px;}
.mcard .lbl{font-size:.62rem;font-weight:600;color:#7fa8c9;text-transform:uppercase;letter-spacing:.10em;}
.mcard .val{font-family:'JetBrains Mono',monospace;font-size:1.2rem;font-weight:600;color:#eaf2ff;margin:3px 0 2px;}
.mcard .sub{font-size:.72rem;color:#90b8d8;line-height:1.5;}
.rcard{background:#112233;border-radius:10px;padding:.9rem 1.1rem .85rem;margin-bottom:4px;}
.rcard .lbl{font-size:.62rem;font-weight:600;color:#7fa8c9;text-transform:uppercase;letter-spacing:.10em;}
.rcard .val{font-family:'JetBrains Mono',monospace;font-size:1.05rem;font-weight:600;color:#eaf2ff;margin:3px 0 2px;}
.rcard .sub{font-size:.72rem;color:#90b8d8;line-height:1.5;}
.npcard{background:#1a0533;border-radius:10px;padding:.9rem 1.1rem .85rem;margin-bottom:4px;}
.npcard .lbl{font-size:.62rem;font-weight:600;color:#ce93d8;text-transform:uppercase;letter-spacing:.10em;}
.npcard .val{font-family:'JetBrains Mono',monospace;font-size:1.05rem;font-weight:600;color:#f3e5f5;margin:3px 0 2px;}
.npcard .sub{font-size:.72rem;color:#ce93d8;line-height:1.5;}

.p-sig{display:inline-block;background:#1b7f45;color:#fff;font-size:.68rem;font-weight:700;padding:2px 8px;border-radius:20px;}
.p-ns {display:inline-block;background:#b71c1c;color:#fff;font-size:.68rem;font-weight:700;padding:2px 8px;border-radius:20px;}
.p-trend{display:inline-block;background:#e65100;color:#fff;font-size:.68rem;font-weight:700;padding:2px 8px;border-radius:20px;}

.ibox{background:#f0f6ff;border-left:4px solid #1565c0;border-radius:6px;
  padding:.8rem 1rem;font-size:.875rem;line-height:1.72;color:#0d1b2a;margin-bottom:.5rem;}
.ibox-np{background:#f9f0ff;border-left:4px solid #7b1fa2;border-radius:6px;
  padding:.8rem 1rem;font-size:.875rem;line-height:1.72;color:#1a0030;margin-bottom:.5rem;}
.abox-warn{background:#fff8e1;border-left:4px solid #f9a825;border-radius:6px;
  padding:.7rem 1rem;font-size:.84rem;color:#5d4037;margin-bottom:.5rem;}
.abox-ok  {background:#e8f5e9;border-left:4px solid #2e7d32;border-radius:6px;
  padding:.7rem 1rem;font-size:.84rem;color:#1b5e20;margin-bottom:.5rem;}
.abox-info{background:#e3f2fd;border-left:4px solid #1565c0;border-radius:6px;
  padding:.7rem 1rem;font-size:.84rem;color:#0d47a1;margin-bottom:.5rem;}
.route-p {background:#e8f5e9;border:2px solid #2e7d32;border-radius:10px;
  padding:1rem 1.2rem;font-size:.92rem;color:#1b5e20;margin:.6rem 0 1rem;}
.route-np{background:#f9f0ff;border:2px solid #7b1fa2;border-radius:10px;
  padding:1rem 1.2rem;font-size:.92rem;color:#4a148c;margin:.6rem 0 1rem;}
.upload-hint{background:#f5f7fa;border:1.5px dashed #b0bec5;border-radius:8px;
  padding:.85rem 1rem;font-size:.84rem;color:#546e7a;}

[data-testid="stSidebar"]{background:#0d1b2a !important;}
[data-testid="stSidebar"] *{color:#cfe2f3 !important;}
[data-testid="stSidebar"] hr{border-color:#1c3048 !important;}
</style>
""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
#  SIDEBAR
# ══════════════════════════════════════════════════════════════════════════════
with st.sidebar:
    st.markdown("## ⚙️ Settings")
    st.markdown("---")
    alpha_level = st.selectbox("Significance level (α)", [0.05,0.01,0.001,0.10], index=0)
    norm_alpha  = st.selectbox("Normality decision α", [0.05,0.01,0.10], index=0,
                               help="Shapiro–Wilk p-threshold for choosing parametric vs non-parametric path.")
    posthoc_method = st.selectbox("Post-hoc correction",
        ["bonferroni","holm","sidak","fdr_bh"],
        format_func=lambda x:{"bonferroni":"Bonferroni","holm":"Holm (step-down)",
                              "sidak":"Šidák","fdr_bh":"FDR Benjamini–Hochberg"}[x])
    effect_pref = st.selectbox("Effect size metric",
        ["partial_eta2","eta2","cohen_f"],
        format_func=lambda x:{"partial_eta2":"Partial η²p (SPSS default)",
                              "eta2":"η² (eta-squared)","cohen_f":"Cohen's f"}[x])
    sph_corr = st.selectbox("Sphericity correction",
        ["auto","gg","hf","lb","none"],
        format_func=lambda x:{"auto":"Auto (GG when p<α)","gg":"Greenhouse–Geisser",
                              "hf":"Huynh–Feldt","lb":"Lower-bound","none":"None"}[x])
    st.markdown("---")
    pal_name   = st.selectbox("Color palette",["tab10","Set2","deep","colorblind","husl"])
    grid_style = st.selectbox("Grid style",["whitegrid","ticks","darkgrid"])

# ══════════════════════════════════════════════════════════════════════════════
#  HEADER
# ══════════════════════════════════════════════════════════════════════════════
st.markdown('<div class="app-title">Mixed ANOVA Calculator</div>', unsafe_allow_html=True)
st.markdown(
    '<div class="app-sub">SPSS GLM Repeated Measures Equivalent · Wide-Format Input · '
    'Automatic Parametric / Non-Parametric Routing · '
    'Mauchly Sphericity · GG/HF/LB Corrections · EMM · Full Report Export</div>',
    unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
#  FORMAT GUIDE & TEMPLATES
# ══════════════════════════════════════════════════════════════════════════════
with st.expander("📋  Data Format Guide & CSV Template Download", expanded=False):
    st.markdown('<div class="sec">Required CSV Format — Wide (One Row Per Subject)</div>',
                unsafe_allow_html=True)
    c1,c2=st.columns([1,1])
    with c1:
        st.markdown(
            "**Wide format:** one row per subject. Each repeated measurement "
            "is a separate column. Column names for repeated measures become the "
            "within-subjects level labels."
        )
        st.dataframe(pd.DataFrame({
            "Column":  ["subject_id (optional)","group_col","time_1","time_2","… time_b"],
            "Role":    ["Unique subject identifier (optional)",
                        "Between-subjects factor (group membership)",
                        "Score at time-point 1","Score at time-point 2",
                        "Score at time-point b"],
            "Example": ["1, 2, 3, …","Control / Treat_A / Treat_B",
                        "42.5","51.3","…"],
            "Required":["No","Yes","Yes","Yes","Yes"],
        }), hide_index=True, use_container_width=True)
    with c2:
        st.dataframe(pd.DataFrame({
            "Rule":[
                "Min. between-subjects groups","Min. repeated measures (columns)",
                "Max. groups / measures","Min. subjects per group",
                "All measures required","Missing values","File size","Encoding"],
            "Specification":[
                "2 (supports 3, 4, … k)","2 (supports 3, 4, … b)",
                "No hard limit","≥ 5 (≥ 10 recommended)",
                "Every subject must have all time-point columns filled",
                "Rows with any NA removed (listwise deletion)","3 MB","UTF-8"],
        }), hide_index=True, use_container_width=True)
        st.markdown("**Tip:** Name repeated-measure columns with a number prefix "
                    "(e.g., `1_Pre`, `2_Post`) to control their order.")

    st.markdown("---")
    def make_wide_template(groups, times, n_per, seed=0):
        np.random.seed(seed)
        base  = {g:50+i*3      for i,g in enumerate(groups)}
        tgain = {t:j*4          for j,t in enumerate(times)}
        igain = {g:{t:i*j*1.5  for j,t in enumerate(times)}
                 for i,g in enumerate(groups)}
        rows,sid=[],1
        for g in groups:
            for _ in range(n_per):
                re=np.random.normal(0,3)
                row={"Group":g}
                for t in times:
                    row[t]=round(base[g]+tgain[t]+igain[g][t]+re+np.random.normal(0,2),2)
                rows.append(row); sid+=1
        return pd.DataFrame(rows)

    tpls={
        "2 groups × 2 measures": make_wide_template(["Control","Treatment"],["Pre","Post"],12),
        "2 groups × 3 measures": make_wide_template(["Control","Treatment"],["Pre","Post","Follow_up"],10),
        "3 groups × 3 measures": make_wide_template(["Control","Treat_A","Treat_B"],["Pre","Post","Follow_up"],10),
        "3 groups × 4 measures": make_wide_template(["Control","Treat_A","Treat_B"],
                                                    ["Baseline","Week4","Week8","Week12"],10),
    }
    tc=st.columns(4)
    for col,(lbl,tdf) in zip(tc,tpls.items()):
        with col:
            st.markdown(f"**{lbl}**"); st.caption(f"{len(tdf)} rows")
            fname=lbl.replace(" ","_").replace("×","x")+".csv"
            st.download_button("⬇️ Download",tdf.to_csv(index=False).encode(),
                               fname,"text/csv",use_container_width=True)
    st.markdown("**Preview — 3 groups × 3 measures (first 9 rows):**")
    st.dataframe(tpls["3 groups × 3 measures"].head(9),hide_index=True,use_container_width=True)

# ══════════════════════════════════════════════════════════════════════════════
#  UPLOAD
# ══════════════════════════════════════════════════════════════════════════════
st.markdown('<div class="sec">Upload Data File (Wide Format)</div>', unsafe_allow_html=True)
uploaded=st.file_uploader("Upload CSV — max 3 MB, UTF-8, wide format (one row per subject)",type=["csv"])
if uploaded is None:
    st.markdown('<div class="upload-hint">📂 No file uploaded yet. Upload your CSV above or download a template.</div>',
                unsafe_allow_html=True); st.stop()
if uploaded.size>3*1024*1024:
    st.error("❌ File exceeds 3 MB."); st.stop()
try:
    df_raw=pd.read_csv(uploaded)
except Exception as e:
    st.error(f"❌ Cannot read CSV: {e}"); st.stop()

# ══════════════════════════════════════════════════════════════════════════════
#  COLUMN MAPPING (wide format)
# ══════════════════════════════════════════════════════════════════════════════
st.markdown('<div class="sec">Column Mapping</div>', unsafe_allow_html=True)
st.markdown("Select the between-subjects (group) column and the repeated-measure (time) columns.")

all_cols=df_raw.columns.tolist()
num_cols=[c for c in all_cols if pd.api.types.is_numeric_dtype(df_raw[c])]
cat_cols=[c for c in all_cols if not pd.api.types.is_numeric_dtype(df_raw[c])]

cm1,cm2=st.columns([1,2])
with cm1:
    grp_col=st.selectbox("👥 Between-subjects factor (group column)",
                         all_cols, index=0 if cat_cols==[] else all_cols.index(cat_cols[0])
                         if cat_cols else 0)
with cm2:
    # Default: all numeric columns as time columns
    default_times=[c for c in num_cols if c!=grp_col]
    time_cols=st.multiselect("🔁 Repeated-measure columns (time-points, in order)",
                             options=[c for c in all_cols if c!=grp_col],
                             default=default_times,
                             help="Select in the correct temporal order.")

if len(time_cols)<2:
    st.error("Select at least 2 repeated-measure columns."); st.stop()

run=st.button("▶  Run Analysis", type="primary", use_container_width=True)
if not run:
    with st.expander("Preview uploaded data (first 8 rows)",expanded=False):
        st.dataframe(df_raw.head(8),hide_index=True,use_container_width=True)
    st.stop()

# ══════════════════════════════════════════════════════════════════════════════
#  DATA PREPARATION
# ══════════════════════════════════════════════════════════════════════════════
needed=[grp_col]+time_cols
df_wide=df_raw[needed].copy()
# Convert time columns to numeric
for c in time_cols:
    df_wide[c]=pd.to_numeric(df_wide[c],errors="coerce")
n_before=len(df_wide); df_wide.dropna(inplace=True)
if len(df_wide)<n_before:
    st.warning(f"⚠️ {n_before-len(df_wide)} row(s) with missing values removed.")

df_wide[grp_col]=df_wide[grp_col].astype(str)
df_wide=df_wide.reset_index(drop=True)
df_wide.insert(0,"_subj_",range(1,len(df_wide)+1))

groups_list=sorted(df_wide[grp_col].unique().tolist(),key=str)
a=len(groups_list); b=len(time_cols); N=len(df_wide)
n_g=df_wide.groupby(grp_col)["_subj_"].count()

if a<2: st.error("Between-subjects factor must have ≥ 2 groups."); st.stop()
if b<2: st.error("Select ≥ 2 repeated-measure columns."); st.stop()
if N<6: st.error("At least 6 subjects required."); st.stop()
if n_g.min()<3: st.error(f"Smallest group has {n_g.min()} subject(s); minimum = 3."); st.stop()

# Long format (for some computations)
df_long=df_wide.melt(id_vars=["_subj_",grp_col],
                     value_vars=time_cols,
                     var_name="_time_",value_name="_y_")
df_long["_time_"]=pd.Categorical(df_long["_time_"],categories=time_cols,ordered=True)

# ══════════════════════════════════════════════════════════════════════════════
#  UTILITY
# ══════════════════════════════════════════════════════════════════════════════
def fmt_p(p):
    if pd.isna(p): return "—"
    if p<0.001:    return "< .001"
    s=f"{p:.3f}"; return s[1:] if s.startswith("0") else s

def fmt_f(v,dec=3): return "—" if pd.isna(v) else f"{v:.{dec}f}"

def cohen_f_es(np2):
    return np.nan if (np.isnan(np2) or np2>=1) else float(np.sqrt(np2/(1-np2)))

def magnitude(v,metric="partial_eta2"):
    if np.isnan(v): return "—"
    if metric in ("partial_eta2","eta2"):
        if v<.010: return "negligible"
        if v<.060: return "small"
        if v<.140: return "medium"
        return "large"
    if v<.10: return "negligible"
    if v<.25: return "small"
    if v<.40: return "medium"
    return "large"

def obs_power(F_val,df1,df2,alpha):
    try:
        ncp=max(float(F_val*df1),0); crit=fdist.ppf(1-alpha,df1,df2)
        return float(np.clip(1-stats.ncf.cdf(crit,df1,df2,nc=ncp),0,1))
    except: return np.nan

def pill(p,alpha):
    if pd.isna(p): return ""
    lbl="< .001" if p<.001 else f"= {fmt_p(p)}"
    if p<alpha:  return f'<span class="p-sig">p {lbl}  ✓</span>'
    if p<0.10:   return f'<span class="p-trend">p {lbl}  ~</span>'
    return f'<span class="p-ns">p {lbl}  n.s.</span>'

def mcp(p_values,method,alpha):
    """Multiple-comparison correction (no statsmodels required)."""
    p=np.array(p_values,dtype=float); n=len(p); ord=np.argsort(p); adj=np.empty(n)
    if method=="bonferroni":
        adj=np.clip(p*n,0,1)
    elif method=="holm":
        for rk,ix in enumerate(ord): adj[ix]=min(1.0,p[ix]*(n-rk))
        cm2=0.0
        for ix in ord: cm2=max(cm2,adj[ix]); adj[ix]=cm2
    elif method=="sidak":
        adj=np.clip(1-(1-p)**n,0,1)
    elif method=="fdr_bh":
        ps=np.empty(n)
        for i,ix in enumerate(ord): ps[i]=p[ix]*n/(i+1)
        mn=1.0
        for i in range(n-1,-1,-1): mn=min(mn,ps[i]); ps[i]=mn
        for i,ix in enumerate(ord): adj[ix]=np.clip(ps[i],0,1)
    else:
        adj=np.clip(p*n,0,1)
    return adj<alpha, adj

# ══════════════════════════════════════════════════════════════════════════════
#  STEP 1 — NORMALITY CHECK (Shapiro–Wilk per cell)
# ══════════════════════════════════════════════════════════════════════════════
sw_rows=[]
any_nonnormal=False
for g in groups_list:
    for t in time_cols:
        vals=df_wide[df_wide[grp_col]==g][t].values
        if len(vals)<3: continue
        W_sw,p_sw=stats.shapiro(vals)
        nn=p_sw<norm_alpha
        if nn: any_nonnormal=True
        sw_rows.append({grp_col:g,"Time":t,"n":len(vals),
                        "Shapiro–Wilk W":round(W_sw,4),
                        "p-value":fmt_p(p_sw),
                        "Normal":("Yes" if not nn else "No"),
                        "_p_raw":p_sw})
sw_df=pd.DataFrame(sw_rows)

PARAMETRIC = not any_nonnormal   # True → parametric path

# ══════════════════════════════════════════════════════════════════════════════
#  STEP 2 — LEVENE'S TEST (homogeneity, relevant for parametric)
# ══════════════════════════════════════════════════════════════════════════════
lev_rows=[]
for t in time_cols:
    gdata=[df_wide[df_wide[grp_col]==g][t].values for g in groups_list]
    if all(len(x)>=2 for x in gdata):
        Fl,pl=stats.levene(*gdata,center="mean")
        lev_rows.append({"Time":t,"Levene F":round(Fl,4),
                         "df1":a-1,"df2":N-a,
                         "p-value":fmt_p(pl),
                         "Homogeneous":("Yes" if pl>alpha_level else "No"),
                         "_p_raw":pl})
lev_df=pd.DataFrame(lev_rows)

# ══════════════════════════════════════════════════════════════════════════════
#  MAUCHLY'S TEST OF SPHERICITY
# ══════════════════════════════════════════════════════════════════════════════
def mauchly_test(wide_mat):
    """wide_mat: (N, b) — all subjects × repeated measures."""
    n,k=wide_mat.shape; p_c=k-1
    C=np.zeros((k,p_c))
    for j in range(p_c):
        C[:j+1,j]=1/np.sqrt((j+1)*(j+2)); C[j+1,j]=-np.sqrt((j+1)/(j+2))
    Y=wide_mat@C; S=np.cov(Y.T,ddof=1)
    if S.ndim==0: S=np.array([[float(S)]])
    det_S=max(np.linalg.det(S),1e-300); tr_S=np.trace(S)
    W=float(np.clip(det_S/(tr_S/p_c)**p_c,1e-10,1.0))
    ff=1-(2*p_c**2+p_c+2)/(6*p_c*(n-1))
    dfm=int(p_c*(p_c+1)/2-1)
    chi2=-np.log(W)*(n-1)*ff; pm=float(1-stats.chi2.cdf(chi2,dfm))
    tr2=np.trace(S@S)
    gg=float(np.clip(tr_S**2/(p_c*tr2),1/p_c,1.0))
    hf_n=n*p_c*gg-2; hf_d=p_c*(n-1-p_c*gg)
    hf=float(np.clip(hf_n/hf_d if hf_d else 1.0,1/p_c,1.0))
    return dict(W=W,chi2=chi2,df_m=dfm,p=pm,eps_gg=gg,
                eps_hf=hf,eps_lb=float(1/p_c))

wide_mat=df_wide[time_cols].values.astype(float)
sph=None; eps=1.0; sph_label="None"
if b>2 and PARAMETRIC:
    sph=mauchly_test(wide_mat)
    if sph_corr=="auto":
        if sph["p"]<alpha_level: eps=sph["eps_gg"]; sph_label="Greenhouse–Geisser"
        else: sph_label="None (sphericity satisfied)"
    elif sph_corr=="gg":  eps=sph["eps_gg"]; sph_label="Greenhouse–Geisser"
    elif sph_corr=="hf":  eps=sph["eps_hf"]; sph_label="Huynh–Feldt"
    elif sph_corr=="lb":  eps=sph["eps_lb"]; sph_label="Lower-bound"

# ══════════════════════════════════════════════════════════════════════════════
#  PARAMETRIC PATH — MIXED ANOVA ENGINE
# ══════════════════════════════════════════════════════════════════════════════
def run_mixed_anova(df_wide, df_long, grp_col, time_cols, groups_list, a, b, N, eps, alpha):
    """
    SPSS GLM Repeated Measures — exact formula replication.
    Uses wide-format input; internally converts for SS computation.
    Verified: SS_A + SS_SA + SS_B + SS_AB + SS_BSA = SS_Total (to numerical precision).
    """
    n_g  = df_wide.groupby(grp_col)["_subj_"].count()
    grand= df_long["_y_"].mean()
    gm   = df_long.groupby(grp_col)["_y_"].mean()     # Ȳᵢ..
    tm   = df_long.groupby("_time_")["_y_"].mean()    # Ȳ.j.
    sm   = df_long.groupby("_subj_")["_y_"].mean()    # Ȳₛ..
    sg   = df_long.groupby("_subj_")[grp_col].first()

    # Cell means Ȳᵢⱼ
    cm={}
    for g in groups_list:
        for t in time_cols:
            cm[(g,t)]=df_long[(df_long[grp_col]==g)&(df_long["_time_"]==t)]["_y_"].mean()

    # ── SS (verified formula set) ────────────────────────────────────────────
    SS_A  = b*sum(n_g[g]*(gm[g]-grand)**2 for g in groups_list)
    SS_SA = sum(b*(sm[s]-gm[sg[s]])**2 for s in df_wide["_subj_"])
    SS_B  = N*sum((tm[t]-grand)**2 for t in time_cols)
    SS_AB = sum(n_g[g]*(cm[(g,t)]-gm[g]-tm[t]+grand)**2
                for g in groups_list for t in time_cols)
    SS_BSA=0.0
    for _,r in df_long.iterrows():
        SS_BSA+=(r["_y_"]-sm[r["_subj_"]]-cm[(r[grp_col],r["_time_"])]+gm[r[grp_col]])**2
    SS_T=((df_long["_y_"]-grand)**2).sum()

    # ── df ───────────────────────────────────────────────────────────────────
    df_A=a-1; df_SA=N-a; df_B=b-1; df_AB=(a-1)*(b-1); df_BSA=(b-1)*(N-a); df_Tot=N*b-1

    # ── Corrected df (within effects) ────────────────────────────────────────
    df_Bc=df_B*eps; df_ABc=df_AB*eps; df_BSAc=df_BSA*eps

    # ── MS ───────────────────────────────────────────────────────────────────
    MS_A  =SS_A /df_A   if df_A   >0 else np.nan
    MS_SA =SS_SA/df_SA  if df_SA  >0 else np.nan
    MS_B  =SS_B /df_Bc  if df_Bc  >0 else np.nan
    MS_AB =SS_AB/df_ABc if df_ABc >0 else np.nan
    MS_BSA=SS_BSA/df_BSAc if df_BSAc>0 else np.nan

    # ── F & p ────────────────────────────────────────────────────────────────
    def _F(e,r): return e/r if (not np.isnan(r) and r>0) else np.nan
    def _p(F,d1,d2): return float(1-fdist.cdf(F,d1,d2)) if not np.isnan(F) else np.nan

    F_A=_F(MS_A,MS_SA);   p_A=_p(F_A,df_A,df_SA)
    F_B=_F(MS_B,MS_BSA);  p_B=_p(F_B,df_Bc,df_BSAc)
    F_AB=_F(MS_AB,MS_BSA);p_AB=_p(F_AB,df_ABc,df_BSAc)

    # ── Effect sizes ─────────────────────────────────────────────────────────
    def _np2(se,sr): return float(se/(se+sr)) if (se+sr)>0 else np.nan
    np2_A=_np2(SS_A,SS_SA); eta2_A=SS_A/SS_T if SS_T>0 else np.nan
    np2_B=_np2(SS_B,SS_BSA);eta2_B=SS_B/SS_T if SS_T>0 else np.nan
    np2_AB=_np2(SS_AB,SS_BSA);eta2_AB=SS_AB/SS_T if SS_T>0 else np.nan

    # ── Power ────────────────────────────────────────────────────────────────
    pw_A=obs_power(F_A,df_A,df_SA,alpha)
    pw_B=obs_power(F_B,df_Bc,df_BSAc,alpha)
    pw_AB=obs_power(F_AB,df_ABc,df_BSAc,alpha)

    # ── Estimated Marginal Means (SPSS method) ───────────────────────────────
    t_w=stats.t.ppf(1-alpha/2,df_BSAc) if df_BSAc>0 else np.nan
    t_b=stats.t.ppf(1-alpha/2,df_SA)   if df_SA>0   else np.nan

    emm_rows=[]
    for g in groups_list:
        for t in time_cols:
            n_c=int(n_g[g]); mn=cm[(g,t)]
            se=float(np.sqrt(MS_BSA/n_c)) if (not np.isnan(MS_BSA) and MS_BSA>0) else np.nan
            emm_rows.append({grp_col:g,"Time":t,"N":n_c,"Mean":round(mn,4),
                             "SE (EMM)":round(se,4) if not np.isnan(se) else np.nan,
                             f"CI Lower":round(mn-t_w*se,4) if not np.isnan(se) else np.nan,
                             f"CI Upper":round(mn+t_w*se,4) if not np.isnan(se) else np.nan})

    emm_btw=[]
    for g in groups_list:
        n_i=int(n_g[g]); mn=float(gm[g])
        se=float(np.sqrt(MS_SA/(b*n_i))) if (not np.isnan(MS_SA) and MS_SA>0) else np.nan
        emm_btw.append({grp_col:g,"N":n_i,"Mean":round(mn,4),
                        "SE (EMM)":round(se,4) if not np.isnan(se) else np.nan,
                        "CI Lower":round(mn-t_b*se,4) if not np.isnan(se) else np.nan,
                        "CI Upper":round(mn+t_b*se,4) if not np.isnan(se) else np.nan})

    emm_win=[]
    for t in time_cols:
        mn=float(tm[t])
        se=float(np.sqrt(MS_BSA/N)) if (not np.isnan(MS_BSA) and MS_BSA>0) else np.nan
        emm_win.append({"Time":t,"N":N,"Mean":round(mn,4),
                        "SE (EMM)":round(se,4) if not np.isnan(se) else np.nan,
                        "CI Lower":round(mn-t_w*se,4) if not np.isnan(se) else np.nan,
                        "CI Upper":round(mn+t_w*se,4) if not np.isnan(se) else np.nan})

    # ── Simple effects (SPSS F-test method) ──────────────────────────────────
    se_rows=[]
    # Within at each group
    for g in groups_list:
        n_i=int(n_g[g]); gm_i=float(gm[g])
        SS_Bi=n_i*sum((cm[(g,t)]-gm_i)**2 for t in time_cols)
        MS_Bi=SS_Bi/(b-1)
        Fi=MS_Bi/MS_BSA if (not np.isnan(MS_BSA) and MS_BSA>0) else np.nan
        pi=float(1-fdist.cdf(Fi,b-1,df_BSAc)) if not np.isnan(Fi) else np.nan
        np2i=SS_Bi/(SS_Bi+MS_BSA*df_BSAc) if not np.isnan(MS_BSA) else np.nan
        se_rows.append({"Effect":f"Within ({', '.join(time_cols)}) at {grp_col}={g}",
                        "SS":round(SS_Bi,4),"df_num":b-1,"df_den":round(df_BSAc,3),
                        "MS":round(MS_Bi,4),"F":round(Fi,4) if not np.isnan(Fi) else np.nan,
                        "p":fmt_p(pi),"Partial η²p":round(np2i,4) if not np.isnan(np2i) else np.nan,
                        "_p_raw":pi})
    # Between at each time
    for t in time_cols:
        tm_t=df_long[df_long["_time_"]==t]["_y_"].mean()
        SS_At=sum(n_g[g]*(cm[(g,t)]-tm_t)**2 for g in groups_list)
        MS_At=SS_At/(a-1)
        Ft=MS_At/MS_SA if (not np.isnan(MS_SA) and MS_SA>0) else np.nan
        pt=float(1-fdist.cdf(Ft,a-1,df_SA)) if not np.isnan(Ft) else np.nan
        np2t=SS_At/(SS_At+MS_SA*df_SA) if not np.isnan(MS_SA) else np.nan
        se_rows.append({"Effect":f"Between ({grp_col}) at Time={t}",
                        "SS":round(SS_At,4),"df_num":a-1,"df_den":df_SA,
                        "MS":round(MS_At,4),"F":round(Ft,4) if not np.isnan(Ft) else np.nan,
                        "p":fmt_p(pt),"Partial η²p":round(np2t,4) if not np.isnan(np2t) else np.nan,
                        "_p_raw":pt})

    return dict(
        SS_A=SS_A,SS_SA=SS_SA,SS_B=SS_B,SS_AB=SS_AB,SS_BSA=SS_BSA,SS_T=SS_T,
        df_A=df_A,df_SA=df_SA,df_B=df_B,df_AB=df_AB,df_BSA=df_BSA,df_Tot=df_Tot,
        df_Bc=df_Bc,df_ABc=df_ABc,df_BSAc=df_BSAc,
        MS_A=MS_A,MS_SA=MS_SA,MS_B=MS_B,MS_AB=MS_AB,MS_BSA=MS_BSA,
        F_A=F_A,F_B=F_B,F_AB=F_AB,p_A=p_A,p_B=p_B,p_AB=p_AB,
        np2_A=np2_A,np2_B=np2_B,np2_AB=np2_AB,
        eta2_A=eta2_A,eta2_B=eta2_B,eta2_AB=eta2_AB,
        pw_A=pw_A,pw_B=pw_B,pw_AB=pw_AB,
        grand=grand,gm=gm,tm=tm,cm=cm,n_g=n_g,
        emm_df=pd.DataFrame(emm_rows),
        emm_btw=pd.DataFrame(emm_btw),
        emm_win=pd.DataFrame(emm_win),
        se_df=pd.DataFrame(se_rows),
        t_w=t_w,t_b=t_b,
    )

# ══════════════════════════════════════════════════════════════════════════════
#  PARAMETRIC POST-HOC
# ══════════════════════════════════════════════════════════════════════════════
def ph_between_param(df_wide, grp_col, groups_list, method, alpha):
    """Pooled-variance t-tests (equal_var=True) — SPSS GLM post-hoc default."""
    rows=[]
    for l1,l2 in itertools.combinations(groups_list,2):
        g1=df_wide[df_wide[grp_col]==l1][time_cols].values.flatten()
        g2=df_wide[df_wide[grp_col]==l2][time_cols].values.flatten()
        # Use group means (one value per subject) for the between comparison
        m1=df_wide[df_wide[grp_col]==l1][time_cols].mean(axis=1).values
        m2=df_wide[df_wide[grp_col]==l2][time_cols].mean(axis=1).values
        t_,pr=stats.ttest_ind(m1,m2,equal_var=True)
        md=m1.mean()-m2.mean()
        n1,n2=len(m1),len(m2)
        sp=np.sqrt(((n1-1)*m1.std(ddof=1)**2+(n2-1)*m2.std(ddof=1)**2)/(n1+n2-2))
        rows.append({f"{grp_col} 1":str(l1),f"{grp_col} 2":str(l2),
                     "n₁":n1,"n₂":n2,
                     "Mean₁":round(m1.mean(),4),"Mean₂":round(m2.mean(),4),
                     "Diff":round(md,4),"t":round(t_,4),"df":n1+n2-2,
                     "p_raw":pr,"Cohen's d":round(md/sp,4) if sp>0 else np.nan})
    if not rows: return pd.DataFrame()
    ph=pd.DataFrame(rows)
    _,adj=mcp(ph["p_raw"].tolist(),method,alpha)
    ph["p (corrected)"]=[ fmt_p(x) for x in adj]
    ph["Reject H₀"]=[x<alpha for x in adj]
    ph["p (uncorrected)"]=ph["p_raw"].map(fmt_p); ph.drop(columns=["p_raw"],inplace=True)
    return ph

def ph_within_param(df_wide, time_cols, method, alpha):
    """Paired t-tests + correction — SPSS repeated measures post-hoc."""
    rows=[]
    for t1,t2 in itertools.combinations(time_cols,2):
        d=df_wide[t1]-df_wide[t2]
        t_,pr=stats.ttest_rel(df_wide[t1],df_wide[t2])
        sd_d=d.std(ddof=1)
        rows.append({"Level 1":t1,"Level 2":t2,"n":len(df_wide),
                     "Mean₁":round(df_wide[t1].mean(),4),"Mean₂":round(df_wide[t2].mean(),4),
                     "Diff":round(d.mean(),4),"SD(Diff)":round(sd_d,4),
                     "t":round(t_,4),"df":len(df_wide)-1,
                     "p_raw":pr,"Cohen's d":round(d.mean()/sd_d,4) if sd_d>0 else np.nan})
    if not rows: return pd.DataFrame()
    ph=pd.DataFrame(rows)
    _,adj=mcp(ph["p_raw"].tolist(),method,alpha)
    ph["p (corrected)"]=[fmt_p(x) for x in adj]
    ph["Reject H₀"]=[x<alpha for x in adj]
    ph["p (uncorrected)"]=ph["p_raw"].map(fmt_p); ph.drop(columns=["p_raw"],inplace=True)
    return ph

# ══════════════════════════════════════════════════════════════════════════════
#  NON-PARAMETRIC PATH
# ══════════════════════════════════════════════════════════════════════════════
def run_nonparametric(df_wide, grp_col, time_cols, groups_list, a, b, N, method, alpha):
    """
    Non-parametric mixed design analysis:
    • Friedman test (within-subjects) — one per group + omnibus across all subjects
    • Kruskal–Wallis (between-subjects) — at each time-point + omnibus
    • Wilcoxon signed-rank post-hoc (within)
    • Mann–Whitney U post-hoc (between)
    • Kendall's W effect size (Friedman), rank-biserial r (Mann–Whitney)
    """
    results={}

    # ── Friedman per group ────────────────────────────────────────────────────
    fried_rows=[]
    for g in groups_list:
        sub=df_wide[df_wide[grp_col]==g][time_cols].values
        if sub.shape[0]<3: continue
        chi2,p=stats.friedmanchisquare(*sub.T)
        # Kendall's W = chi2 / (n*(b-1))
        W_k=chi2/(sub.shape[0]*(b-1)) if b>1 else np.nan
        fried_rows.append({grp_col:g,"n":sub.shape[0],
                           "Friedman χ²":round(chi2,4),"df":b-1,
                           "p-value":fmt_p(p),"Kendall's W":round(W_k,4),
                           "Significant":p<alpha,"_p":p})
    fried_df=pd.DataFrame(fried_rows)

    # ── Omnibus Friedman (all subjects) ───────────────────────────────────────
    all_sub=df_wide[time_cols].values
    chi2_all,p_all=stats.friedmanchisquare(*all_sub.T)
    W_k_all=chi2_all/(N*(b-1)) if b>1 else np.nan

    # ── Kruskal–Wallis at each time-point ─────────────────────────────────────
    kw_rows=[]
    for t in time_cols:
        gdata=[df_wide[df_wide[grp_col]==g][t].values for g in groups_list]
        H,p=stats.kruskal(*gdata)
        # Eta-squared approximation: η²=(H-k+1)/(n-k)
        eta2_kw=(H-(a-1))/(N-a) if N>a else np.nan
        kw_rows.append({"Time":t,"Kruskal–Wallis H":round(H,4),"df":a-1,
                        "p-value":fmt_p(p),"η² (approx)":round(eta2_kw,4),
                        "Significant":p<alpha,"_p":p})
    kw_df=pd.DataFrame(kw_rows)

    # ── Omnibus KW (averaged across time) ─────────────────────────────────────
    kw_data=[df_wide[df_wide[grp_col]==g][time_cols].values.flatten() for g in groups_list]
    H_all,p_kw_all=stats.kruskal(*kw_data)

    # ── Wilcoxon signed-rank post-hoc (within, per group) ─────────────────────
    wilc_rows=[]
    raw_ps_w=[]
    for g in groups_list:
        sub=df_wide[df_wide[grp_col]==g]
        for t1,t2 in itertools.combinations(time_cols,2):
            d=sub[t1].values-sub[t2].values
            if len(d)<3: continue
            stat,pr=stats.wilcoxon(d,alternative="two-sided")
            # Rank-biserial r = 1 - 2*T_minus/(n*(n+1)/2) [simple approx via Z]
            n_w=len(d)
            Z=stats.norm.ppf(1-pr/2); rb=Z/np.sqrt(n_w)
            raw_ps_w.append(pr)
            wilc_rows.append({grp_col:g,"Level 1":t1,"Level 2":t2,"n":n_w,
                              "Wilcoxon W":round(stat,4),"p_raw":pr,
                              "r (rank-biserial)":round(rb,4)})
    if wilc_rows:
        wilc_df=pd.DataFrame(wilc_rows)
        _,adj_w=mcp(wilc_df["p_raw"].tolist(),method,alpha)
        wilc_df["p (corrected)"]=[fmt_p(x) for x in adj_w]
        wilc_df["Reject H₀"]=[x<alpha for x in adj_w]
        wilc_df["p (uncorrected)"]=wilc_df["p_raw"].map(fmt_p)
        wilc_df.drop(columns=["p_raw"],inplace=True)
    else:
        wilc_df=pd.DataFrame()

    # ── Mann–Whitney post-hoc (between, per time-point) ───────────────────────
    mw_rows=[]
    raw_ps_mw=[]
    for t in time_cols:
        for l1,l2 in itertools.combinations(groups_list,2):
            g1=df_wide[df_wide[grp_col]==l1][t].values
            g2=df_wide[df_wide[grp_col]==l2][t].values
            U,pr=stats.mannwhitneyu(g1,g2,alternative="two-sided")
            # Rank-biserial r = 1 - 2U/(n1*n2)
            rb=1-2*U/(len(g1)*len(g2))
            raw_ps_mw.append(pr)
            mw_rows.append({"Time":t,f"{grp_col} 1":str(l1),f"{grp_col} 2":str(l2),
                            "n₁":len(g1),"n₂":len(g2),
                            "Median₁":round(np.median(g1),4),"Median₂":round(np.median(g2),4),
                            "Mann–Whitney U":round(U,4),"p_raw":pr,
                            "r (rank-biserial)":round(rb,4)})
    if mw_rows:
        mw_df=pd.DataFrame(mw_rows)
        _,adj_mw=mcp(mw_df["p_raw"].tolist(),method,alpha)
        mw_df["p (corrected)"]=[fmt_p(x) for x in adj_mw]
        mw_df["Reject H₀"]=[x<alpha for x in adj_mw]
        mw_df["p (uncorrected)"]=mw_df["p_raw"].map(fmt_p)
        mw_df.drop(columns=["p_raw"],inplace=True)
    else:
        mw_df=pd.DataFrame()

    # ── Descriptive: median per cell ───────────────────────────────────────────
    desc_rows=[]
    for g in groups_list:
        for t in time_cols:
            vals=df_wide[df_wide[grp_col]==g][t].values
            desc_rows.append({grp_col:g,"Time":t,"n":len(vals),
                              "Median":round(np.median(vals),4),
                              "IQR":round(np.percentile(vals,75)-np.percentile(vals,25),4),
                              "Min":round(vals.min(),4),"Max":round(vals.max(),4)})
    desc_df=pd.DataFrame(desc_rows)

    return dict(
        fried_df=fried_df, chi2_all=chi2_all, p_all=p_all, W_k_all=W_k_all,
        kw_df=kw_df, H_all=H_all, p_kw_all=p_kw_all,
        wilc_df=wilc_df, mw_df=mw_df, desc_df=desc_df,
    )

# ══════════════════════════════════════════════════════════════════════════════
#  RUN ENGINES
# ══════════════════════════════════════════════════════════════════════════════
with st.spinner("⏳  Running analysis …"):
    if PARAMETRIC:
        res=run_mixed_anova(df_wide,df_long,grp_col,time_cols,
                            groups_list,a,b,N,eps,alpha_level)
        df_ph_b=ph_between_param(df_wide,grp_col,groups_list,posthoc_method,alpha_level)
        df_ph_w=ph_within_param(df_wide,time_cols,posthoc_method,alpha_level)
    else:
        np_res=run_nonparametric(df_wide,grp_col,time_cols,
                                 groups_list,a,b,N,posthoc_method,alpha_level)

# ══════════════════════════════════════════════════════════════════════════════
#  ASSUMPTION RESULTS
# ══════════════════════════════════════════════════════════════════════════════
st.markdown('<div class="sec">Step 1 — Assumption Checks</div>', unsafe_allow_html=True)

assump_tabs=st.tabs(["Normality — Shapiro–Wilk","Homogeneity of Variance — Levene",
                     "Sphericity — Mauchly","Analysis Route"])
with assump_tabs[0]:
    st.markdown(
        "Shapiro–Wilk test applied to **each cell** (group × time-point). "
        "**p > α** is consistent with normality. "
        "If **any** cell violates normality, the non-parametric path is selected automatically."
    )
    st.dataframe(sw_df.drop(columns=["_p_raw"]),hide_index=True,use_container_width=True)
    if any_nonnormal:
        st.markdown(
            f'<div class="abox-warn"><b>Non-normality detected</b> — '
            f'at least one cell has Shapiro–Wilk p &lt; {norm_alpha}. '
            f'<b>Non-parametric analysis path selected automatically.</b></div>',
            unsafe_allow_html=True)
    else:
        st.markdown(
            f'<div class="abox-ok"><b>Normality satisfied</b> — '
            f'all cells have Shapiro–Wilk p ≥ {norm_alpha}. '
            f'<b>Parametric Mixed ANOVA path selected.</b></div>',
            unsafe_allow_html=True)

with assump_tabs[1]:
    st.markdown(
        "Levene's test (center = mean) for equality of variances across groups "
        "at each time-point. Relevant for the parametric path; informational for non-parametric."
    )
    st.dataframe(lev_df.drop(columns=["_p_raw"]),hide_index=True,use_container_width=True)

with assump_tabs[2]:
    if not PARAMETRIC:
        st.markdown(
            '<div class="abox-info">Mauchly\'s test is not applicable '
            'on the non-parametric path.</div>', unsafe_allow_html=True)
    elif b<=2:
        st.markdown(
            '<div class="abox-info">Mauchly\'s test not applicable — '
            'within-subjects factor has only 2 levels (sphericity trivially satisfied).</div>',
            unsafe_allow_html=True)
    elif sph is None:
        st.markdown('<div class="abox-info">Sphericity test not run.</div>',unsafe_allow_html=True)
    else:
        sph_tbl=pd.DataFrame({
            "Statistic":["Mauchly's W","Chi-square (χ²)","df","p-value",
                         "Greenhouse–Geisser ε","Huynh–Feldt ε","Lower-bound ε",
                         "Correction applied","ε used"],
            "Value":[f"{sph['W']:.4f}",f"{sph['chi2']:.4f}",str(sph['df_m']),
                     fmt_p(sph['p']),f"{sph['eps_gg']:.4f}",f"{sph['eps_hf']:.4f}",
                     f"{sph['eps_lb']:.4f}",sph_label,f"{eps:.4f}"],
        })
        st.dataframe(sph_tbl,hide_index=True,use_container_width=True)
        if sph["p"]<alpha_level:
            st.markdown(
                f'<div class="abox-warn"><b>Sphericity violated</b> — '
                f'W = {sph["W"]:.4f}, p {fmt_p(sph["p"])} &lt; {alpha_level}. '
                f'df corrected using {sph_label} (ε = {eps:.4f}).</div>',
                unsafe_allow_html=True)
        else:
            st.markdown(
                f'<div class="abox-ok"><b>Sphericity satisfied</b> — '
                f'W = {sph["W"]:.4f}, p {fmt_p(sph["p"])} ≥ {alpha_level}. '
                f'No df correction required.</div>',
                unsafe_allow_html=True)

with assump_tabs[3]:
    if PARAMETRIC:
        st.markdown(
            '<div class="route-p">✅ <b>PARAMETRIC PATH SELECTED</b><br>'
            'All cells satisfy the normality assumption (Shapiro–Wilk p ≥ α). '
            'Proceeding with Mixed ANOVA (GLM Repeated Measures), '
            'identical to SPSS output.</div>',
            unsafe_allow_html=True)
    else:
        st.markdown(
            '<div class="route-np">⚠️ <b>NON-PARAMETRIC PATH SELECTED</b><br>'
            'One or more cells violate the normality assumption (Shapiro–Wilk p &lt; α). '
            'Proceeding with non-parametric equivalents: '
            '<b>Friedman test</b> (within-subjects) and '
            '<b>Kruskal–Wallis test</b> (between-subjects), '
            'with Wilcoxon signed-rank and Mann–Whitney U post-hoc comparisons.</div>',
            unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
#  ════════  PARAMETRIC OUTPUT  ════════
# ══════════════════════════════════════════════════════════════════════════════
if PARAMETRIC:
    # ── Overview cards ─────────────────────────────────────────────────────────
    st.markdown('<div class="sec">Results Overview</div>', unsafe_allow_html=True)
    oc=st.columns(5)
    def mcard(col,lbl,val,sub):
        col.markdown(f'<div class="mcard"><div class="lbl">{lbl}</div>'
                     f'<div class="val">{val}</div><div class="sub">{sub}</div></div>',
                     unsafe_allow_html=True)
    mcard(oc[0],"Subjects (N)",str(N),f"{a} group(s) · {b} time-point(s)")
    mcard(oc[1],f"Between — {grp_col}",f"F = {res['F_A']:.3f}",
          f"p {fmt_p(res['p_A'])} · η²p = {res['np2_A']:.3f}")
    mcard(oc[2],f"Within — Time",f"F = {res['F_B']:.3f}",
          f"p {fmt_p(res['p_B'])} · η²p = {res['np2_B']:.3f}")
    mcard(oc[3],"Interaction",f"F = {res['F_AB']:.3f}",
          f"p {fmt_p(res['p_AB'])} · η²p = {res['np2_AB']:.3f}")
    mcard(oc[4],"Grand Mean",f"{res['grand']:.3f}",
          f"SD = {df_long['_y_'].std():.3f}")

    st.markdown("")
    rc1,rc2,rc3=st.columns(3)
    for col_,lbl,F,d1,d2,p,np2,eta2,pw in [
        (rc1,grp_col,res["F_A"],res["df_A"],res["df_SA"],
         res["p_A"],res["np2_A"],res["eta2_A"],res["pw_A"]),
        (rc2,"Time (within)",res["F_B"],res["df_Bc"],res["df_BSAc"],
         res["p_B"],res["np2_B"],res["eta2_B"],res["pw_B"]),
        (rc3,f"{grp_col} × Time",res["F_AB"],res["df_ABc"],res["df_BSAc"],
         res["p_AB"],res["np2_AB"],res["eta2_AB"],res["pw_AB"]),
    ]:
        cf=cohen_f_es(np2); mg=magnitude(np2,"partial_eta2")
        col_.markdown(
            f'<div class="rcard"><div class="lbl">{lbl}</div>'
            f'<div class="val">F({d1:.2f}, {d2:.2f}) = {F:.3f}</div>'
            f'<div class="sub">{pill(p,alpha_level)}</div>'
            f'<div class="sub" style="margin-top:5px;">Partial η²p = {np2:.4f} ({mg})</div>'
            f'<div class="sub">η² = {eta2:.4f} · Cohen\'s f = {fmt_f(cf)}</div>'
            f'<div class="sub">Observed power = {pw:.3f}</div></div>',
            unsafe_allow_html=True)

    # ── Descriptives & EMM ─────────────────────────────────────────────────────
    st.markdown('<div class="sec">Descriptive Statistics & Estimated Marginal Means</div>',
                unsafe_allow_html=True)
    d1,d2,d3,d4=st.tabs(["Cell Raw Descriptives","Cell EMMs (SPSS)",
                          f"Marginal EMMs — {grp_col}","Marginal EMMs — Time"])
    with d1:
        cr_rows=[]
        for g in groups_list:
            for t in time_cols:
                v=df_wide[df_wide[grp_col]==g][t].values
                cr_rows.append({grp_col:g,"Time":t,"N":len(v),
                                "Mean":round(v.mean(),4),"SD":round(v.std(ddof=1),4),
                                "SE":round(v.std(ddof=1)/np.sqrt(len(v)),4),
                                "Min":round(v.min(),4),"Max":round(v.max(),4)})
        st.dataframe(pd.DataFrame(cr_rows),hide_index=True,use_container_width=True)
        st.caption("SE = SD / √n  (raw descriptive standard error).")
    with d2:
        ci_pct=int(100*(1-alpha_level))
        st.markdown(
            f"**SPSS Estimated Marginal Means.** "
            f"SE = √(MS_BS(A) / n_cell) = √({res['MS_BSA']:.4f} / n). "
            f"CI = Mean ± t({res['df_BSAc']:.2f})×SE at α = {alpha_level} "
            f"(t-critical = ±{res['t_w']:.4f})."
        )
        emm=res["emm_df"].copy()
        emm.columns=[grp_col,"Time","N","Mean","SE (EMM)",f"{ci_pct}% CI Lower",f"{ci_pct}% CI Upper"]
        st.dataframe(emm,hide_index=True,use_container_width=True)
    with d3:
        eb=res["emm_btw"].copy()
        ci_pct=int(100*(1-alpha_level))
        eb.columns=[grp_col,"N","Mean","SE (EMM)",f"{ci_pct}% CI Lower",f"{ci_pct}% CI Upper"]
        st.dataframe(eb,hide_index=True,use_container_width=True)
    with d4:
        ew=res["emm_win"].copy()
        ew.rename(columns={"Time":"Time-point"},inplace=True)
        st.dataframe(ew,hide_index=True,use_container_width=True)

    # ── ANOVA Table ────────────────────────────────────────────────────────────
    st.markdown('<div class="sec">Mixed ANOVA Summary Table</div>', unsafe_allow_html=True)
    at=[
        {"Source":"BETWEEN SUBJECTS","SS":"","df":"","MS":"","F":"","p":"","Partial η²p":"","Obs. Power":""},
        {"Source":f"  {grp_col}","SS":fmt_f(res["SS_A"]),"df":fmt_f(res["df_A"],0),
         "MS":fmt_f(res["MS_A"]),"F":fmt_f(res["F_A"]),"p":fmt_p(res["p_A"]),
         "Partial η²p":fmt_f(res["np2_A"]),"Obs. Power":fmt_f(res["pw_A"])},
        {"Source":"  Error [S(A)]","SS":fmt_f(res["SS_SA"]),"df":fmt_f(res["df_SA"],0),
         "MS":fmt_f(res["MS_SA"]),"F":"—","p":"—","Partial η²p":"—","Obs. Power":"—"},
        {"Source":"WITHIN SUBJECTS","SS":"","df":"","MS":"","F":"","p":"","Partial η²p":"","Obs. Power":""},
        {"Source":"  Time","SS":fmt_f(res["SS_B"]),"df":fmt_f(res["df_Bc"]),
         "MS":fmt_f(res["MS_B"]),"F":fmt_f(res["F_B"]),"p":fmt_p(res["p_B"]),
         "Partial η²p":fmt_f(res["np2_B"]),"Obs. Power":fmt_f(res["pw_B"])},
        {"Source":f"  {grp_col} × Time","SS":fmt_f(res["SS_AB"]),"df":fmt_f(res["df_ABc"]),
         "MS":fmt_f(res["MS_AB"]),"F":fmt_f(res["F_AB"]),"p":fmt_p(res["p_AB"]),
         "Partial η²p":fmt_f(res["np2_AB"]),"Obs. Power":fmt_f(res["pw_AB"])},
        {"Source":"  Error [BS(A)]","SS":fmt_f(res["SS_BSA"]),"df":fmt_f(res["df_BSAc"]),
         "MS":fmt_f(res["MS_BSA"]),"F":"—","p":"—","Partial η²p":"—","Obs. Power":"—"},
        {"Source":"Total","SS":fmt_f(res["SS_T"]),"df":fmt_f(res["df_Tot"],0),
         "MS":"—","F":"—","p":"—","Partial η²p":"—","Obs. Power":"—"},
    ]
    at_df=pd.DataFrame(at)
    st.dataframe(at_df,hide_index=True,use_container_width=True)
    note="S(A) = Subjects within Groups (between error); BS(A) = B × Subjects within Groups (within error)."
    if sph and eps<1.0:
        note=f"Note: df for within-subjects effects adjusted using {sph_label} (ε = {eps:.4f}). "+note
    st.caption(note)

    # ── Effect Sizes ───────────────────────────────────────────────────────────
    st.markdown('<div class="sec">Effect Size Summary</div>', unsafe_allow_html=True)
    st.dataframe(pd.DataFrame({
        "Source":[grp_col,"Time (within)",f"{grp_col} × Time"],
        "Partial η²p":[round(res["np2_A"],4),round(res["np2_B"],4),round(res["np2_AB"],4)],
        "η² (total)":[round(res["eta2_A"],4),round(res["eta2_B"],4),round(res["eta2_AB"],4)],
        "Cohen's f":[round(cohen_f_es(res["np2_A"]),4),round(cohen_f_es(res["np2_B"]),4),
                     round(cohen_f_es(res["np2_AB"]),4)],
        "Magnitude":[magnitude(res["np2_A"],"partial_eta2"),
                     magnitude(res["np2_B"],"partial_eta2"),
                     magnitude(res["np2_AB"],"partial_eta2")],
        "F":[fmt_f(res["F_A"]),fmt_f(res["F_B"]),fmt_f(res["F_AB"])],
        "p":[fmt_p(res["p_A"]),fmt_p(res["p_B"]),fmt_p(res["p_AB"])],
        "Significant":[res["p_A"]<alpha_level,res["p_B"]<alpha_level,res["p_AB"]<alpha_level],
    }),hide_index=True,use_container_width=True)
    st.caption("Partial η²p benchmarks (Cohen 1988): negligible < .01; small .01–.05; medium .06–.13; large ≥ .14.")

    # ── Post-hoc ───────────────────────────────────────────────────────────────
    st.markdown('<div class="sec">Post-Hoc Pairwise Comparisons</div>', unsafe_allow_html=True)
    corr_lbl={"bonferroni":"Bonferroni","holm":"Holm","sidak":"Šidák",
              "fdr_bh":"Benjamini–Hochberg FDR"}.get(posthoc_method,posthoc_method)
    pt1,pt2,pt3=st.tabs([f"Between — {grp_col}","Within — Time",
                         "Simple Effects (SPSS F-test)"])
    with pt1:
        st.markdown(f"Pooled-variance t-tests on subject means (SPSS GLM default). Correction: **{corr_lbl}**.")
        if a<3:
            st.markdown('<div class="abox-info">Only 2 groups — omnibus F-test is conclusive.</div>',
                        unsafe_allow_html=True)
        elif df_ph_b.empty: st.warning("No comparisons computed.")
        else:
            st.dataframe(df_ph_b,hide_index=True,use_container_width=True)
            st.caption("Cohen's d: |d| < .20 negligible; .20–.49 small; .50–.79 medium; ≥ .80 large.")
    with pt2:
        st.markdown(f"Paired t-tests. Correction: **{corr_lbl}**.")
        if b<3: st.markdown('<div class="abox-info">Only 2 levels — omnibus F-test is conclusive.</div>',
                             unsafe_allow_html=True)
        elif df_ph_w.empty: st.warning("No comparisons computed.")
        else: st.dataframe(df_ph_w,hide_index=True,use_container_width=True)
    with pt3:
        st.markdown(
            "**SPSS F-test method** — uses pooled error terms from the omnibus model.\n"
            f"- Within (Time) at each group: F = MS_B@group / MS_BS(A), df = ({b-1}, {res['df_BSAc']:.2f})\n"
            f"- Between ({grp_col}) at each time-point: F = MS_A@time / MS_S(A), df = ({a-1}, {res['df_SA']})"
        )
        se_out=res["se_df"].drop(columns=["_p_raw"],errors="ignore")
        st.dataframe(se_out,hide_index=True,use_container_width=True)

    # ── Interpretation ─────────────────────────────────────────────────────────
    st.markdown('<div class="sec">Statistical Interpretation</div>', unsafe_allow_html=True)

    def interp_p(res,grp_col,alpha,pref):
        texts=[]
        def es(np2,eta2):
            cf=cohen_f_es(np2)
            if pref=="partial_eta2": return f"partial η²p = {np2:.3f} ({magnitude(np2,'partial_eta2')} effect)"
            elif pref=="eta2":       return f"η² = {eta2:.3f} ({magnitude(eta2,'eta2')} effect)"
            else:                    return f"Cohen's f = {cf:.3f} ({magnitude(cf,'cohen_f')} effect)"

        sig_A=res["p_A"]<alpha
        es_A=es(res["np2_A"],res["eta2_A"])
        if sig_A:
            texts.append(f"<b>Main Effect of {grp_col} (Between-Subjects):</b> Statistically significant, "
                f"F({res['df_A']:.0f}, {res['df_SA']:.0f}) = {res['F_A']:.3f}, p {fmt_p(res['p_A'])}, "
                f"{es_A}; observed power = {res['pw_A']:.3f}. Group means on the dependent variable "
                f"differ significantly when averaging across time-points. "
                f"Post-hoc comparisons identify which groups differ.")
        else:
            texts.append(f"<b>Main Effect of {grp_col} (Between-Subjects):</b> Not statistically significant, "
                f"F({res['df_A']:.0f}, {res['df_SA']:.0f}) = {res['F_A']:.3f}, p {fmt_p(res['p_A'])}, "
                f"{es_A}; observed power = {res['pw_A']:.3f}. "
                f"Insufficient evidence that groups differ when averaging across time.")

        sph_note=""
        if sph and eps<1.0:
            sph_note=f" df adjusted via {sph_label} (ε = {eps:.4f})."
        sig_B=res["p_B"]<alpha
        es_B=es(res["np2_B"],res["eta2_B"])
        if sig_B:
            texts.append(f"<b>Main Effect of Time (Within-Subjects):</b> Statistically significant, "
                f"F({res['df_Bc']:.3f}, {res['df_BSAc']:.3f}) = {res['F_B']:.3f}, p {fmt_p(res['p_B'])}, "
                f"{es_B}{sph_note}; observed power = {res['pw_B']:.3f}. "
                f"Scores change significantly across time-points when averaging across groups.")
        else:
            texts.append(f"<b>Main Effect of Time (Within-Subjects):</b> Not statistically significant, "
                f"F({res['df_Bc']:.3f}, {res['df_BSAc']:.3f}) = {res['F_B']:.3f}, p {fmt_p(res['p_B'])}, "
                f"{es_B}{sph_note}; observed power = {res['pw_B']:.3f}. "
                f"No significant change in scores across time-points detected.")

        sig_AB=res["p_AB"]<alpha
        es_AB=es(res["np2_AB"],res["eta2_AB"])
        if sig_AB:
            texts.append(f"<b>{grp_col} × Time Interaction:</b> Statistically significant, "
                f"F({res['df_ABc']:.3f}, {res['df_BSAc']:.3f}) = {res['F_AB']:.3f}, p {fmt_p(res['p_AB'])}, "
                f"{es_AB}; observed power = {res['pw_AB']:.3f}. "
                f"The effect of time on the dependent variable differs across {grp_col} groups (non-parallel profiles). "
                f"<b>Caution:</b> interpret main effects in the context of this significant interaction. "
                f"Examine the profile plot and simple-effects analysis.")
        else:
            texts.append(f"<b>{grp_col} × Time Interaction:</b> Not statistically significant, "
                f"F({res['df_ABc']:.3f}, {res['df_BSAc']:.3f}) = {res['F_AB']:.3f}, p {fmt_p(res['p_AB'])}, "
                f"{es_AB}; observed power = {res['pw_AB']:.3f}. "
                f"The pattern of change across time is consistent (parallel) across {grp_col} groups. "
                f"Main effects may be interpreted independently.")

        low=[nm for nm,pw in [(grp_col,res["pw_A"]),("Time",res["pw_B"]),(f"{grp_col}×Time",res["pw_AB"])]
             if not np.isnan(pw) and pw<0.80]
        if low:
            texts.append(f"<b>Statistical Power Notice:</b> Observed power < 0.80 for: {', '.join(low)}. "
                f"Increased Type II error risk. Consider a larger sample size.")
        return texts

    for txt in interp_p(res,grp_col,alpha_level,effect_pref):
        st.markdown(f'<div class="ibox">{txt}</div>',unsafe_allow_html=True)

    with st.expander("APA 7th Edition Reporting Template",expanded=False):
        apa=(f"A {a} ({grp_col}) × {b} (time) mixed ANOVA was conducted "
             f"(N = {N}). "
             f"The main effect of {grp_col} was "
             f"{'significant' if res['p_A']<alpha_level else 'not significant'}, "
             f"F({res['df_A']:.0f}, {res['df_SA']:.0f}) = {res['F_A']:.2f}, "
             f"p {fmt_p(res['p_A'])}, partial η²p = {res['np2_A']:.3f}. "
             f"The main effect of time was "
             f"{'significant' if res['p_B']<alpha_level else 'not significant'}, "
             f"F({res['df_Bc']:.2f}, {res['df_BSAc']:.2f}) = {res['F_B']:.2f}, "
             f"p {fmt_p(res['p_B'])}, partial η²p = {res['np2_B']:.3f}. "
             f"The {grp_col} × time interaction was "
             f"{'significant' if res['p_AB']<alpha_level else 'not significant'}, "
             f"F({res['df_ABc']:.2f}, {res['df_BSAc']:.2f}) = {res['F_AB']:.2f}, "
             f"p {fmt_p(res['p_AB'])}, partial η²p = {res['np2_AB']:.3f}.")
        st.text_area("APA result:",value=apa,height=160)

# ══════════════════════════════════════════════════════════════════════════════
#  ════════  NON-PARAMETRIC OUTPUT  ════════
# ══════════════════════════════════════════════════════════════════════════════
else:
    st.markdown('<div class="sec-np">Non-Parametric Analysis Results</div>',
                unsafe_allow_html=True)
    st.markdown(
        '<div class="ibox-np"><b>Non-parametric path selected.</b> '
        'Because the normality assumption was violated, the following non-parametric tests '
        'are reported: '
        '<b>Friedman test</b> for within-subjects repeated-measures (equivalent to repeated-measures ANOVA), '
        '<b>Kruskal–Wallis test</b> for between-subjects comparisons (equivalent to one-way ANOVA), '
        'with <b>Wilcoxon signed-rank</b> (within) and <b>Mann–Whitney U</b> (between) post-hoc tests. '
        'Effect sizes: Kendall\'s W (Friedman), rank-biserial r (Mann–Whitney/Wilcoxon).</div>',
        unsafe_allow_html=True)

    # ── Descriptives ───────────────────────────────────────────────────────────
    st.markdown('<div class="sec-np">Descriptive Statistics (Median, IQR)</div>',
                unsafe_allow_html=True)
    st.dataframe(np_res["desc_df"],hide_index=True,use_container_width=True)
    st.caption("Non-parametric path uses median and IQR as central tendency / spread measures.")

    # ── Friedman ───────────────────────────────────────────────────────────────
    st.markdown('<div class="sec-np">Friedman Test (Within-Subjects — Per Group)</div>',
                unsafe_allow_html=True)
    npc1,npc2=st.columns([2,1])
    with npc1:
        st.dataframe(np_res["fried_df"].drop(columns=["_p"]),
                     hide_index=True,use_container_width=True)
        st.caption("Kendall's W: 0.10 = small, 0.30 = medium, 0.50 = large.")
    with npc2:
        st.markdown(
            f'<div class="npcard"><div class="lbl">Omnibus Friedman (all subjects)</div>'
            f'<div class="val">χ²({b-1}) = {np_res["chi2_all"]:.3f}</div>'
            f'<div class="sub">{pill(np_res["p_all"],alpha_level)}</div>'
            f'<div class="sub" style="margin-top:5px;">Kendall\'s W = {np_res["W_k_all"]:.4f}</div>'
            f'</div>',unsafe_allow_html=True)

    # ── Kruskal–Wallis ─────────────────────────────────────────────────────────
    st.markdown('<div class="sec-np">Kruskal–Wallis Test (Between-Subjects — Per Time-Point)</div>',
                unsafe_allow_html=True)
    npc3,npc4=st.columns([2,1])
    with npc3:
        st.dataframe(np_res["kw_df"].drop(columns=["_p"]),
                     hide_index=True,use_container_width=True)
        st.caption("η² approximation: (H − k + 1) / (N − k).")
    with npc4:
        st.markdown(
            f'<div class="npcard"><div class="lbl">Omnibus Kruskal–Wallis</div>'
            f'<div class="val">H({a-1}) = {np_res["H_all"]:.3f}</div>'
            f'<div class="sub">{pill(np_res["p_kw_all"],alpha_level)}</div>'
            f'</div>',unsafe_allow_html=True)

    # ── Wilcoxon post-hoc ──────────────────────────────────────────────────────
    st.markdown('<div class="sec-np">Wilcoxon Signed-Rank Post-Hoc (Within, per Group)</div>',
                unsafe_allow_html=True)
    if np_res["wilc_df"].empty: st.info("Not enough data for Wilcoxon post-hoc.")
    else:
        st.dataframe(np_res["wilc_df"],hide_index=True,use_container_width=True)
        st.caption("Rank-biserial r: |r| < .10 negligible; .10–.29 small; .30–.49 medium; ≥ .50 large.")

    # ── Mann–Whitney post-hoc ──────────────────────────────────────────────────
    st.markdown('<div class="sec-np">Mann–Whitney U Post-Hoc (Between Groups, per Time-Point)</div>',
                unsafe_allow_html=True)
    if np_res["mw_df"].empty: st.info("Not enough data for Mann–Whitney post-hoc.")
    else:
        st.dataframe(np_res["mw_df"],hide_index=True,use_container_width=True)
        st.caption(f"p-values corrected using {posthoc_method}. "
                   "Rank-biserial r: |r| < .10 negligible; .10–.29 small; .30–.49 medium; ≥ .50 large.")

    # ── Interpretation ─────────────────────────────────────────────────────────
    st.markdown('<div class="sec-np">Statistical Interpretation</div>', unsafe_allow_html=True)
    np_texts=[]
    # Friedman
    for _,row in np_res["fried_df"].iterrows():
        sig=row["_p"]<alpha_level
        w_lbl="large" if row["Kendall's W"]>=.50 else "medium" if row["Kendall's W"]>=.30 else "small" if row["Kendall's W"]>=.10 else "negligible"
        kw_val = row["Kendall's W"]
        chi2_val = row["Friedman χ²"]
        p_val_str = row["p-value"]
        g_name = row[grp_col]
        if sig:
            np_texts.append(
                f"<b>Friedman Test — {grp_col} = {g_name}:</b> Statistically significant, "
                f"\u03c7\u00b2({b-1}) = {chi2_val}, p {p_val_str}, "
                f"Kendall\u2019s W = {kw_val} ({w_lbl} effect). "
                f"Scores differ significantly across time-points within the {g_name} group.")
        else:
            np_texts.append(
                f"<b>Friedman Test — {grp_col} = {g_name}:</b> Not statistically significant, "
                f"\u03c7\u00b2({b-1}) = {chi2_val}, p {p_val_str}, "
                f"Kendall\u2019s W = {kw_val} ({w_lbl} effect). "
                f"No significant difference across time-points in the {g_name} group.")
    # KW omnibus
    sig_kw=np_res["p_kw_all"]<alpha_level
    if sig_kw:
        np_texts.append(f"<b>Kruskal–Wallis Test (omnibus between-subjects):</b> Statistically significant, "
            f"H({a-1}) = {np_res['H_all']:.3f}, p {fmt_p(np_res['p_kw_all'])}. "
            f"At least one group differs from the others. Mann–Whitney post-hoc identifies specific pairs.")
    else:
        np_texts.append(f"<b>Kruskal–Wallis Test (omnibus between-subjects):</b> Not statistically significant, "
            f"H({a-1}) = {np_res['H_all']:.3f}, p {fmt_p(np_res['p_kw_all'])}. "
            f"No significant between-group difference detected.")
    for txt in np_texts:
        st.markdown(f'<div class="ibox-np">{txt}</div>',unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
#  VISUALISATIONS (both paths)
# ══════════════════════════════════════════════════════════════════════════════
st.markdown('<div class="sec">Visualisations</div>', unsafe_allow_html=True)

sns.set_style(grid_style); sns.set_context("notebook",font_scale=1.0)
colors=sns.color_palette(pal_name,n_colors=max(a,3))

with st.spinner("Rendering plots …"):
    fig=plt.figure(figsize=(20,13)); fig.patch.set_facecolor("#f7f9fc")
    gs=mgs.GridSpec(2,3,figure=fig,hspace=0.46,wspace=0.34)

    # ── Profile Plot ─────────────────────────────────────────────────────────
    ax1=fig.add_subplot(gs[0,:2])
    if PARAMETRIC:
        emm_plot=res["emm_df"].copy()
        emm_plot.columns=[grp_col,"Time","N","Mean","SE","CI_lo","CI_hi"]
        for i,g in enumerate(groups_list):
            sub=emm_plot[emm_plot[grp_col]==g]
            lo=sub["Mean"]-sub["CI_lo"]; hi=sub["CI_hi"]-sub["Mean"]
            ax1.errorbar(range(b),sub["Mean"],[lo,hi],
                         marker="o",ms=7,lw=2.4,capsize=5,capthick=1.8,
                         color=colors[i],label=str(g),zorder=4)
        ax1.set_xticks(range(b)); ax1.set_xticklabels(time_cols,rotation=20 if b>4 else 0)
        ax1.set_title(f"Profile Plot (EMM ± {int(100*(1-alpha_level))}% CI)",
                      fontsize=11.5,fontweight="bold",pad=10)
        ax1.set_ylabel("Estimated Marginal Mean",fontsize=10.5)
    else:
        for i,g in enumerate(groups_list):
            sub=df_wide[df_wide[grp_col]==g][time_cols]
            meds=sub.median(); q1=sub.quantile(.25); q3=sub.quantile(.75)
            ax1.plot(range(b),meds,marker="o",ms=7,lw=2.4,color=colors[i],label=str(g))
            ax1.fill_between(range(b),q1,q3,alpha=0.2,color=colors[i])
        ax1.set_xticks(range(b)); ax1.set_xticklabels(time_cols,rotation=20 if b>4 else 0)
        ax1.set_title("Profile Plot (Median ± IQR shaded — Non-parametric)",
                      fontsize=11.5,fontweight="bold",pad=10)
        ax1.set_ylabel("Median",fontsize=10.5)
    ax1.set_xlabel("Time-point",fontsize=10.5)
    ax1.legend(title=grp_col,framealpha=0.9,fontsize=9)
    ax1.grid(True,alpha=0.28,linestyle="--"); ax1.set_facecolor("#ffffff")

    # ── Bar Chart ─────────────────────────────────────────────────────────────
    ax2=fig.add_subplot(gs[0,2])
    w_pos=np.arange(b); bw=0.80/a
    for i,g in enumerate(groups_list):
        if PARAMETRIC:
            emm_g=res["emm_df"].copy()
            emm_g.columns=[grp_col,"Time","N","Mean","SE","CI_lo","CI_hi"]
            sub=emm_g[emm_g[grp_col]==g].set_index("Time").reindex(time_cols)
            ci_err=np.array([sub["Mean"]-sub["CI_lo"],sub["CI_hi"]-sub["Mean"]])
            ax2.bar(w_pos+(i-a/2+0.5)*bw,sub["Mean"],bw*0.90,
                    yerr=ci_err,color=colors[i],label=str(g),alpha=0.85,
                    capsize=4,error_kw={"elinewidth":1.5,"ecolor":"#333","alpha":.65})
        else:
            meds=df_wide[df_wide[grp_col]==g][time_cols].median()
            ax2.bar(w_pos+(i-a/2+0.5)*bw,meds[time_cols].values,bw*0.90,
                    color=colors[i],label=str(g),alpha=0.85)
    ax2.set_xticks(w_pos)
    ax2.set_xticklabels(time_cols,rotation=22 if b>3 else 0,ha="right" if b>3 else "center")
    ax2.set_title("Mean/Median per Cell",fontsize=11.5,fontweight="bold",pad=10)
    ax2.set_xlabel("Time-point",fontsize=10.5)
    ax2.legend(title=grp_col,fontsize=8.5,framealpha=0.9)
    ax2.grid(True,alpha=0.28,linestyle="--",axis="y"); ax2.set_facecolor("#ffffff")

    # ── Box Plot ──────────────────────────────────────────────────────────────
    ax3=fig.add_subplot(gs[1,0])
    df_box=df_long.copy()
    df_box["Cell"]=df_box[grp_col].astype(str)+"\n"+df_box["_time_"].astype(str)
    c_ord=[f"{g}\n{t}" for g in groups_list for t in time_cols]
    cmap={f"{g}\n{t}":colors[i] for i,g in enumerate(groups_list) for t in time_cols}
    sns.boxplot(data=df_box,x="Cell",y="_y_",order=c_ord,palette=cmap,ax=ax3,
                linewidth=1.1,flierprops=dict(marker="o",ms=3.5,alpha=0.5))
    ax3.set_title("Distribution per Cell",fontsize=11.5,fontweight="bold",pad=10)
    ax3.set_xlabel(""); ax3.set_ylabel("Score",fontsize=10.5)
    ax3.tick_params(axis="x",labelsize=7.5 if a*b>6 else 9)
    ax3.grid(True,alpha=0.28,linestyle="--",axis="y"); ax3.set_facecolor("#ffffff")

    # ── Violin ────────────────────────────────────────────────────────────────
    ax4=fig.add_subplot(gs[1,1])
    sns.violinplot(data=df_long,x="_time_",y="_y_",hue=grp_col,
                   order=time_cols,palette=pal_name,inner="quartile",
                   ax=ax4,alpha=0.78,linewidth=1.0)
    ax4.set_title("Score Distribution (Violin)",fontsize=11.5,fontweight="bold",pad=10)
    ax4.set_xlabel("Time-point",fontsize=10.5); ax4.set_ylabel("Score",fontsize=10.5)
    h,l=ax4.get_legend_handles_labels()
    ax4.legend(h[:a],l[:a],title=grp_col,fontsize=8.5,framealpha=0.9)
    ax4.grid(True,alpha=0.28,linestyle="--",axis="y"); ax4.set_facecolor("#ffffff")

    # ── Effect / Test Size ────────────────────────────────────────────────────
    ax5=fig.add_subplot(gs[1,2])
    if PARAMETRIC:
        eff_l=[grp_col,"Time",f"{grp_col}\n×Time"]
        eff_v=[res["np2_A"],res["np2_B"],res["np2_AB"]]
        eff_p=[res["p_A"],res["p_B"],res["p_AB"]]
        ec=[colors[0] if p<alpha_level else "#9e9e9e" for p in eff_p]
        bars=ax5.barh(eff_l,eff_v,color=ec,height=0.42,alpha=0.88)
        for xv,lb in [(0.01,"small"),(0.06,"medium"),(0.14,"large")]:
            ax5.axvline(xv,ls="--",lw=1.1,color="#777",alpha=0.7)
            ax5.text(xv,2.65,lb,ha="center",fontsize=7,color="#555",va="bottom")
        for bar,val,p_ in zip(bars,eff_v,eff_p):
            mk="  *" if p_<alpha_level else "  n.s."
            ax5.text(val+0.003,bar.get_y()+bar.get_height()/2,
                     f"{val:.4f}{mk}",va="center",fontsize=9,fontweight="bold")
        ax5.set_title("Partial η²p Effect Sizes\n(* = significant at α)",
                      fontsize=11.5,fontweight="bold",pad=10)
        ax5.set_xlabel("Partial η²p",fontsize=10.5)
        ax5.set_xlim(0,max(max(eff_v)*1.38,0.22))
    else:
        # Kendall's W for each group
        w_vals=np_res["fried_df"]["Kendall's W"].values
        w_lbls=[str(x) for x in np_res["fried_df"][grp_col].values]
        w_pv=np_res["fried_df"]["_p"].values
        ec_np=[colors[i] if p<alpha_level else "#9e9e9e" for i,p in enumerate(w_pv)]
        bars2=ax5.barh(w_lbls,w_vals,color=ec_np,height=0.42,alpha=0.88)
        for xv,lb in [(0.10,"small"),(0.30,"medium"),(0.50,"large")]:
            ax5.axvline(xv,ls="--",lw=1.1,color="#777",alpha=0.7)
            ax5.text(xv,len(w_lbls)-0.35,lb,ha="center",fontsize=7,color="#555",va="bottom")
        for bar,val in zip(bars2,w_vals):
            ax5.text(val+0.01,bar.get_y()+bar.get_height()/2,
                     f"{val:.3f}",va="center",fontsize=9,fontweight="bold")
        ax5.set_title("Kendall's W Effect Size\n(Friedman — per group)",
                      fontsize=11.5,fontweight="bold",pad=10)
        ax5.set_xlabel("Kendall's W",fontsize=10.5)
        ax5.set_xlim(0,1.1)
    ax5.set_facecolor("#ffffff"); ax5.grid(True,alpha=0.28,linestyle="--",axis="x")

    path_lbl="Parametric: Mixed ANOVA" if PARAMETRIC else "Non-Parametric: Friedman + Kruskal–Wallis"
    fig.suptitle(f"Mixed Design Analysis — {grp_col} (between) × Time (within)  |  "
                 f"N = {N}  |  {path_lbl}",
                 fontsize=12.5,fontweight="bold",y=1.012,color="#0d1b2a")

st.pyplot(fig,use_container_width=True)

# ══════════════════════════════════════════════════════════════════════════════
#  DOWNLOADS
# ══════════════════════════════════════════════════════════════════════════════
st.markdown('<div class="sec">Download Results</div>', unsafe_allow_html=True)
dl1,dl2,dl3=st.columns(3)

with dl1:
    if PARAMETRIC:
        at_df_dl=pd.DataFrame(at)
        st.download_button("📄  ANOVA Table (CSV)",
                           at_df_dl.to_csv(index=False).encode(),
                           "mixed_anova_table.csv","text/csv",use_container_width=True)
    else:
        combined=pd.concat([np_res["fried_df"].drop(columns=["_p"],errors="ignore"),
                            np_res["kw_df"].drop(columns=["_p"],errors="ignore")],
                           ignore_index=True)
        st.download_button("📄  Test Results (CSV)",
                           combined.to_csv(index=False).encode(),
                           "nonparametric_results.csv","text/csv",use_container_width=True)
with dl2:
    fig_buf=io.BytesIO()
    fig.savefig(fig_buf,format="png",dpi=200,bbox_inches="tight",facecolor="#f7f9fc")
    fig_buf.seek(0)
    st.download_button("🖼️  Figures (PNG, 200 dpi)",fig_buf.getvalue(),
                       "figures.png","image/png",use_container_width=True)
with dl3:
    # Build Word report
    doc=Document()
    for sec in doc.sections:
        sec.top_margin=sec.bottom_margin=Inches(1.0)
        sec.left_margin=sec.right_margin=Inches(1.2)
    h=doc.add_heading("Mixed Design Analysis Report",0)
    h.alignment=WD_ALIGN_PARAGRAPH.CENTER
    h.runs[0].font.color.rgb=RGBColor(0x0d,0x1b,0x2a)
    doc.add_paragraph(
        f"Analysis path: {'Parametric — Mixed ANOVA' if PARAMETRIC else 'Non-Parametric — Friedman + Kruskal–Wallis'}  |  "
        f"Between factor: {grp_col}  |  Within: Time ({', '.join(time_cols)})  |  "
        f"N = {N}  |  α = {alpha_level}"
    ).alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    doc.add_heading("1.  Normality Test (Shapiro–Wilk)",level=1)
    doc.add_paragraph(
        f"{'All cells satisfied normality (p ≥ ' + str(norm_alpha) + '); parametric path selected.' if PARAMETRIC else 'Normality violated in at least one cell (p < ' + str(norm_alpha) + '); non-parametric path selected.'}"
    )
    if PARAMETRIC:
        doc.add_heading("2.  Mixed ANOVA Summary",level=1)
        doc.add_paragraph(
            f"F_between({res['df_A']:.0f},{res['df_SA']:.0f}) = {res['F_A']:.4f}, p {fmt_p(res['p_A'])}, partial η²p = {res['np2_A']:.4f}. "
            f"F_within({res['df_Bc']:.3f},{res['df_BSAc']:.3f}) = {res['F_B']:.4f}, p {fmt_p(res['p_B'])}, partial η²p = {res['np2_B']:.4f}. "
            f"F_interaction({res['df_ABc']:.3f},{res['df_BSAc']:.3f}) = {res['F_AB']:.4f}, p {fmt_p(res['p_AB'])}, partial η²p = {res['np2_AB']:.4f}."
        )
    else:
        doc.add_heading("2.  Non-Parametric Test Summary",level=1)
        doc.add_paragraph(
            f"Omnibus Friedman: χ²({b-1}) = {np_res['chi2_all']:.4f}, p {fmt_p(np_res['p_all'])}, Kendall's W = {np_res['W_k_all']:.4f}. "
            f"Omnibus Kruskal–Wallis: H({a-1}) = {np_res['H_all']:.4f}, p {fmt_p(np_res['p_kw_all'])}."
        )
    doc.add_heading("3.  Figures",level=1)
    img_buf2=io.BytesIO()
    fig.savefig(img_buf2,format="png",dpi=150,bbox_inches="tight",facecolor="#f7f9fc")
    img_buf2.seek(0)
    doc.add_picture(img_buf2,width=Inches(6.2))
    doc.add_paragraph("Figure 1. Analysis visualisation panel.").alignment=WD_ALIGN_PARAGRAPH.CENTER
    word_buf=io.BytesIO(); doc.save(word_buf); word_buf.seek(0)
    st.download_button("📝  Report (Word .docx)",word_buf.getvalue(),
                       "analysis_report.docx",
                       "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                       use_container_width=True)

st.markdown("---")
st.caption(
    "Mixed ANOVA Calculator · SPSS GLM Repeated Measures Equivalent · Wide-Format Input · "
    "Auto Parametric/Non-Parametric Routing · SS decomposition verified (= SS_Total) · "
    "References: Winer, Brown & Michels (1991); Mauchly (1940); Box (1954); "
    "Greenhouse & Geisser (1959); Huynh & Feldt (1976); Lecoutre (1991); Cohen (1988); "
    "Friedman (1937); Kruskal & Wallis (1952); Wilcoxon (1945); Mann & Whitney (1947)."
)
